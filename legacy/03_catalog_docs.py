from pathlib import Path
from datetime import datetime
import subprocess
import tempfile
import csv
import re
import sys

from docx import Document


ROOT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path.cwd()

OUTPUT = ROOT / "raw_text"
OUTPUT.mkdir(exist_ok=True)

TEXT_DIR = OUTPUT / "text"
TEXT_DIR.mkdir(exist_ok=True)


LIBREOFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"


KEYWORDS = [
    "novel",
    "manuscript",
    "chapter",
    "fiction",
    "story",
    "character",
    "protagonist",
    "antagonist",
    "plot",
    "scene",
    "dialogue",
    "narrator",
    "outline",
    "synopsis",
    "draft",
    "writing",
    "author",
    "query",
    "publisher",
    "editor",
    "agent",
    "title",
    "copyright",
]


def extract_docx(path):

    try:

        doc = Document(path)

        parts = []

        for paragraph in doc.paragraphs:

            if paragraph.text.strip():
                parts.append(paragraph.text)

        for table in doc.tables:

            for row in table.rows:

                parts.append(
                    " | ".join(
                        cell.text
                        for cell in row.cells
                    )
                )

        return "\n".join(parts)

    except Exception as e:

        return ""


def convert_doc(path, output_dir):

    try:

        result = subprocess.run(
            [
                LIBREOFFICE,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                str(path),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        converted = output_dir / (
            path.stem + ".docx"
        )

        if converted.exists():
            return converted

    except Exception:
        pass

    return None


def get_keywords(text):

    lower = text.lower()

    found = []

    for keyword in KEYWORDS:

        if re.search(
            rf"\b{re.escape(keyword)}\b",
            lower
        ):
            found.append(keyword)

    return found


# ============================================================
# FIND DOCUMENTS
# ============================================================

docs = [
    p for p in ROOT.rglob("*")
    if p.is_file()
    and p.suffix.lower() in {".doc", ".docx"}
    and "raw_text" not in p.parts
]

print()
print("=" * 60)
print("WORD DOCUMENT CATALOG")
print("=" * 60)
print()
print(f"Found {len(docs):,} Word documents")
print()


records = []


# ============================================================
# PROCESS
# ============================================================

for i, path in enumerate(docs, 1):

    text = ""

    # --------------------------------------------------------
    # DOCX
    # --------------------------------------------------------

    if path.suffix.lower() == ".docx":

        text = extract_docx(path)

    # --------------------------------------------------------
    # DOC
    # --------------------------------------------------------

    else:

        with tempfile.TemporaryDirectory() as temp:

            converted = convert_doc(
                path,
                Path(temp),
            )

            if converted:

                text = extract_docx(converted)

    # --------------------------------------------------------
    # Save extracted text
    # --------------------------------------------------------

    relative = path.relative_to(ROOT)

    safe_name = (
        str(relative)
        .replace("\\", "__")
        .replace("/", "__")
    )

    text_file = TEXT_DIR / (
        safe_name + ".txt"
    )

    text_file.write_text(
        text,
        encoding="utf-8",
        errors="ignore",
    )

    # --------------------------------------------------------
    # Metadata
    # --------------------------------------------------------

    try:

        stat = path.stat()

        modified = datetime.fromtimestamp(
            stat.st_mtime
        ).isoformat(
            sep=" ",
            timespec="seconds",
        )

        size_mb = round(
            stat.st_size / 1024 / 1024,
            2,
        )

    except Exception:

        modified = ""
        size_mb = 0

    words = re.findall(
        r"\b[\w'-]+\b",
        text,
    )

    keywords = get_keywords(text)

    records.append({
        "path": str(relative),
        "filename": path.name,
        "extension": path.suffix.lower(),
        "modified": modified,
        "size_mb": size_mb,
        "word_count": len(words),
        "keyword_count": len(keywords),
        "keywords": ", ".join(keywords),
        "text_file": str(
            text_file.relative_to(ROOT)
        ),
        "preview": re.sub(
            r"\s+",
            " ",
            text[:500],
        ),
    })

    if i % 25 == 0 or i == len(docs):

        print(
            f"Processed {i:,}/{len(docs):,}"
        )


# ============================================================
# SAVE CSV
# ============================================================

csv_file = OUTPUT / "word_documents.csv"

with csv_file.open(
    "w",
    newline="",
    encoding="utf-8-sig",
) as f:

    writer = csv.DictWriter(
        f,
        fieldnames=[
            "path",
            "filename",
            "extension",
            "modified",
            "size_mb",
            "word_count",
            "keyword_count",
            "keywords",
            "text_file",
            "preview",
        ],
    )

    writer.writeheader()
    writer.writerows(records)


# ============================================================
# SORTED POTENTIAL WRITING
# ============================================================

writing = sorted(
    records,
    key=lambda r: (
        r["keyword_count"],
        r["word_count"],
    ),
    reverse=True,
)


writing_file = (
    OUTPUT / "potential_writing.csv"
)

with writing_file.open(
    "w",
    newline="",
    encoding="utf-8-sig",
) as f:

    writer = csv.DictWriter(
        f,
        fieldnames=[
            "path",
            "filename",
            "extension",
            "modified",
            "size_mb",
            "word_count",
            "keyword_count",
            "keywords",
            "text_file",
            "preview",
        ],
    )

    writer.writeheader()

    writer.writerows(
        r for r in writing
        if r["keyword_count"] > 0
    )


print()
print("=" * 60)
print("COMPLETE")
print("=" * 60)
print()
print(f"Catalog:")
print(csv_file)
print()
print(f"Potential writing:")
print(writing_file)
print()
print(f"Extracted text:")
print(TEXT_DIR)