from pathlib import Path
import csv
import subprocess
import tempfile
import re
import sys

try:
    from docx import Document
except ImportError:
    print("Install python-docx first:")
    print("pip install python-docx")
    sys.exit(1)

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None


# ============================================================
# CONFIGURATION
# ============================================================

ROOT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path.cwd()

OUTPUT = ROOT / "_walking_on_air_search"
OUTPUT.mkdir(exist_ok=True)

SEARCH_PHRASES = [
    "walking on air",
]

TEXT_LIMIT = 500_000


# ============================================================
# TEXT EXTRACTION
# ============================================================

def extract_docx(path):
    try:
        doc = Document(path)

        parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                parts.append(
                    " | ".join(
                        cell.text for cell in row.cells
                    )
                )

        return "\n".join(parts)[:TEXT_LIMIT]

    except Exception:
        return ""


def extract_doc(path):
    """
    Convert old .doc file to .docx using LibreOffice,
    then extract the resulting text.
    """

    libreoffice = r"C:\Program Files\LibreOffice\program\soffice.exe"

    with tempfile.TemporaryDirectory() as temp_dir:

        try:
            subprocess.run(
                [
                    libreoffice,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    temp_dir,
                    str(path),
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
            )

            converted = Path(temp_dir) / (
                path.stem + ".docx"
            )

            if converted.exists():
                return extract_docx(converted)

        except Exception:
            pass

    return ""


def extract_pdf(path):

    if PdfReader is None:
        return ""

    try:
        reader = PdfReader(path)

        parts = []

        for page in reader.pages:
            text = page.extract_text()

            if text:
                parts.append(text)

            if len("\n".join(parts)) >= TEXT_LIMIT:
                break

        return "\n".join(parts)[:TEXT_LIMIT]

    except Exception:
        return ""


def extract_text(path):

    ext = path.suffix.lower()

    if ext == ".doc":
        return extract_doc(path)

    if ext == ".docx":
        return extract_docx(path)

    if ext in {".txt", ".rtf", ".md"}:
        try:
            return path.read_text(
                encoding="utf-8",
                errors="ignore",
            )[:TEXT_LIMIT]
        except Exception:
            return ""

    if ext == ".pdf":
        return extract_pdf(path)

    return ""


# ============================================================
# SEARCH
# ============================================================

files = [
    p for p in ROOT.rglob("*")
    if p.is_file()
    and "_walking_on_air_search" not in p.parts
]

print()
print("=" * 60)
print("WALKING ON AIR SEARCH")
print("=" * 60)
print()
print(f"Searching {len(files):,} files...")
print()

results = []

for i, path in enumerate(files, 1):

    # --------------------------------------------------------
    # Filename search
    # --------------------------------------------------------

    filename_matches = []

    filename_lower = path.name.lower()

    for phrase in SEARCH_PHRASES:
        if phrase.lower() in filename_lower:
            filename_matches.append(phrase)

    # --------------------------------------------------------
    # Text extraction
    # --------------------------------------------------------

    text = extract_text(path)

    text_matches = []

    if text:
        text_lower = text.lower()

        for phrase in SEARCH_PHRASES:
            if phrase.lower() in text_lower:
                text_matches.append(phrase)

    # --------------------------------------------------------
    # Record result
    # --------------------------------------------------------

    if filename_matches or text_matches:

        results.append({
            "path": str(path.relative_to(ROOT)),
            "filename": path.name,
            "extension": path.suffix.lower(),
            "filename_match": "; ".join(filename_matches),
            "text_match": "; ".join(text_matches),
            "text_preview": re.sub(
                r"\s+",
                " ",
                text[:1000],
            ),
        })

        print()
        print("FOUND:")
        print(path)

        if filename_matches:
            print("  Filename:", filename_matches)

        if text_matches:
            print("  Text:", text_matches)

    # --------------------------------------------------------
    # Progress
    # --------------------------------------------------------

    if i % 100 == 0 or i == len(files):
        print(
            f"Processed {i:,}/{len(files):,}"
        )


# ============================================================
# SAVE RESULTS
# ============================================================

result_file = OUTPUT / "walking_on_air_results.csv"

with result_file.open(
    "w",
    newline="",
    encoding="utf-8-sig",
) as f:

    writer = csv.DictWriter(
        f,
        fieldnames=[
            "path",
            "filename",
            "extension",
            "filename_match",
            "text_match",
            "text_preview",
        ],
    )

    writer.writeheader()
    writer.writerows(results)


# ============================================================
# SUMMARY
# ============================================================

print()
print("=" * 60)
print("SEARCH COMPLETE")
print("=" * 60)
print()
print(f"Matches found: {len(results):,}")
print()
print(f"Results:")
print(result_file)
print()

if results:

    print("MATCHES:")
    print()

    for result in results:
        print(result["path"])

else:

    print("NO TEXT OR FILENAME MATCHES FOUND.")