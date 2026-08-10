from pathlib import Path
from collections import Counter
import hashlib
import re
import csv
import sys

# Optional dependencies:
# pip install python-docx pypdf openpyxl
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None


# ============================================================
# CONFIGURATION
# ============================================================

ROOT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path.cwd()

OUTPUT = ROOT / "_archive_inventory"
OUTPUT.mkdir(exist_ok=True)

TEXT_LIMIT = 100_000

DOCUMENT_EXTENSIONS = {
    ".txt", ".md", ".rtf",
    ".doc", ".docx",
    ".pdf",
    ".csv",
    ".xlsx", ".xls",
}

AUDIO_EXTENSIONS = {
    ".mp3", ".wav", ".m4a", ".aac", ".flac",
    ".ogg", ".wma", ".aiff",
}

IMAGE_EXTENSIONS = {
    ".jpg", ".jpeg", ".png", ".gif",
    ".tif", ".tiff", ".bmp", ".webp",
}

VIDEO_EXTENSIONS = {
    ".mp4", ".mov", ".avi", ".mkv", ".wmv", ".m4v",
}

# Words/phrases suggesting novel-related material.
KEYWORDS = {
    "walking on air": 10,
    "novel": 5,
    "manuscript": 5,
    "chapter": 5,
    "draft": 4,
    "fiction": 4,
    "story": 4,
    "character": 3,
    "protagonist": 4,
    "antagonist": 4,
    "plot": 3,
    "scene": 3,
    "dialogue": 3,
    "narrator": 3,
    "outline": 3,
    "synopsis": 3,
    "publisher": 2,
    "publishing": 2,
    "editor": 2,
    "agent": 2,
    "book": 2,
    "writing": 3,
    "writer": 3,
    "author": 2,
}

FOLDER_KEYWORDS = {
    "novel": 5,
    "manuscript": 5,
    "writing": 4,
    "draft": 4,
    "fiction": 4,
    "story": 4,
    "book": 3,
    "characters": 3,
    "research": 2,
    "notes": 2,
    "chapter": 4,
}


# ============================================================
# HELPERS
# ============================================================

def classify_file(path):
    ext = path.suffix.lower()

    if ext in DOCUMENT_EXTENSIONS:
        return "document"
    if ext in AUDIO_EXTENSIONS:
        return "audio"
    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in VIDEO_EXTENSIONS:
        return "video"

    return "other"


def file_hash(path, chunk_size=1024 * 1024):
    """SHA-256 hash for exact duplicate detection."""
    h = hashlib.sha256()

    try:
        with path.open("rb") as f:
            while chunk := f.read(chunk_size):
                h.update(chunk)
        return h.hexdigest()
    except (OSError, PermissionError):
        return ""


def extract_text(path):
    """Extract text from common document formats."""

    ext = path.suffix.lower()

    try:
        if ext in {".txt", ".md", ".csv", ".rtf"}:
            return path.read_text(
                encoding="utf-8",
                errors="ignore"
            )[:TEXT_LIMIT]

        if ext == ".docx" and Document:
            doc = Document(path)

            parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    parts.append(
                        " | ".join(cell.text for cell in row.cells)
                    )

            return "\n".join(parts)[:TEXT_LIMIT]

        if ext == ".pdf" and PdfReader:
            reader = PdfReader(path)

            parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)

                if len("\n".join(parts)) >= TEXT_LIMIT:
                    break

            return "\n".join(parts)[:TEXT_LIMIT]

        if ext == ".xlsx" and load_workbook:
            wb = load_workbook(
                path,
                read_only=True,
                data_only=True
            )

            parts = []

            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    values = [
                        str(v) for v in row
                        if v is not None
                    ]

                    if values:
                        parts.append(" | ".join(values))

                    if len("\n".join(parts)) >= TEXT_LIMIT:
                        break

            return "\n".join(parts)[:TEXT_LIMIT]

    except Exception:
        pass

    return ""


def score_relevance(path, text):
    """
    Heuristic relevance score.

    Higher = more likely to contain material related
    to the novel.
    """

    score = 0

    filename = path.stem.lower()
    folder_text = " ".join(
        part.lower()
        for part in path.parent.parts
    )

    combined_name = f"{filename} {folder_text}"

    for keyword, points in KEYWORDS.items():
        if re.search(rf"\b{re.escape(keyword)}\b", filename):
            score += points

    for keyword, points in FOLDER_KEYWORDS.items():
        if re.search(rf"\b{re.escape(keyword)}\b", folder_text):
            score += points

    if text:
        text_lower = text.lower()

        # Don't let repeated occurrences overwhelm the score.
        for keyword, points in KEYWORDS.items():
            occurrences = len(
                re.findall(
                    rf"\b{re.escape(keyword)}\b",
                    text_lower
                )
            )

            score += min(occurrences, 5) * points

        # Prose-like documents get a small bonus.
        words = re.findall(r"\b[A-Za-z]{3,}\b", text)

        if len(words) > 500:
            score += 3

        if len(words) > 2000:
            score += 3

    # Audio is potentially valuable even without text.
    if path.suffix.lower() in AUDIO_EXTENSIONS:
        score += 1

    return score


def relevance_label(score):
    if score >= 15:
        return "HIGH"
    if score >= 7:
        return "POSSIBLE"
    if score >= 3:
        return "LOW"
    return "UNLIKELY"


# ============================================================
# SCAN
# ============================================================

print(f"\nScanning:")
print(f"  {ROOT}\n")

files = [
    p for p in ROOT.rglob("*")
    if p.is_file()
    and "_archive_inventory" not in p.parts
]

print(f"Found {len(files):,} files.\n")

records = []
hashes = {}
extension_counts = Counter()
category_counts = Counter()

for i, path in enumerate(files, 1):

    try:
        stat = path.stat()
    except (OSError, PermissionError):
        continue

    ext = path.suffix.lower() or "[none]"
    category = classify_file(path)

    extension_counts[ext] += 1
    category_counts[category] += 1

    relative_path = path.relative_to(ROOT)

    text = ""

    if category == "document":
        text = extract_text(path)

    score = score_relevance(path, text)
    label = relevance_label(score)

    # Exact duplicate detection.
    file_hash_value = ""

    if stat.st_size > 0:
        file_hash_value = file_hash(path)

        if file_hash_value:
            hashes.setdefault(
                file_hash_value,
                []
            ).append(str(relative_path))

    records.append({
        "path": str(relative_path),
        "filename": path.name,
        "extension": ext,
        "category": category,
        "size_bytes": stat.st_size,
        "size_mb": round(stat.st_size / 1024 / 1024, 2),
        "modified": stat.st_mtime,
        "relevance_score": score,
        "relevance": label,
        "text_extracted": bool(text),
        "text_preview": re.sub(
            r"\s+",
            " ",
            text[:500]
        ),
        "sha256": file_hash_value,
    })

    if i % 100 == 0 or i == len(files):
        print(f"Processed {i:,}/{len(files):,}")


# ============================================================
# DUPLICATES
# ============================================================

duplicate_lookup = {}

for hash_value, paths in hashes.items():
    if len(paths) > 1:
        for path in paths:
            duplicate_lookup[path] = len(paths)

for record in records:
    record["duplicate_count"] = duplicate_lookup.get(
        record["path"],
        1
    )


# ============================================================
# WRITE INVENTORY
# ============================================================

inventory_file = OUTPUT / "file_inventory.csv"

fieldnames = [
    "path",
    "filename",
    "extension",
    "category",
    "size_bytes",
    "size_mb",
    "modified",
    "relevance_score",
    "relevance",
    "text_extracted",
    "text_preview",
    "sha256",
    "duplicate_count",
]

with inventory_file.open(
    "w",
    newline="",
    encoding="utf-8-sig"
) as f:

    writer = csv.DictWriter(
        f,
        fieldnames=fieldnames
    )

    writer.writeheader()
    writer.writerows(records)


# ============================================================
# WRITE NOVEL CANDIDATES
# ============================================================

candidates = sorted(
    records,
    key=lambda r: r["relevance_score"],
    reverse=True
)

candidate_file = OUTPUT / "novel_candidates.csv"

with candidate_file.open(
    "w",
    newline="",
    encoding="utf-8-sig"
) as f:

    writer = csv.DictWriter(
        f,
        fieldnames=fieldnames
    )

    writer.writeheader()

    writer.writerows(
        r for r in candidates
        if r["relevance"] in {"HIGH", "POSSIBLE"}
    )


# ============================================================
# WRITE SUMMARY
# ============================================================

summary_file = OUTPUT / "summary.txt"

with summary_file.open(
    "w",
    encoding="utf-8"
) as f:

    f.write("UNCLE'S FILE ARCHIVE — INITIAL INVENTORY\n")
    f.write("=" * 60 + "\n\n")

    f.write(f"Root directory:\n{ROOT}\n\n")

    f.write(f"Total files: {len(records):,}\n")

    folder_count = len({
        str(Path(r["path"]).parent)
        for r in records
    })
    f.write(f"Total folders: {folder_count:,}\n\n")

    f.write("FILE CATEGORIES\n")
    f.write("-" * 60 + "\n")

    for category, count in category_counts.most_common():
        f.write(f"{category:15} {count:,}\n")

    f.write("\nFILE EXTENSIONS\n")
    f.write("-" * 60 + "\n")

    for ext, count in extension_counts.most_common():
        f.write(f"{ext:15} {count:,}\n")

    f.write("\nRELEVANCE\n")
    f.write("-" * 60 + "\n")

    relevance_counts = Counter(
        r["relevance"]
        for r in records
    )

    for label in ["HIGH", "POSSIBLE", "LOW", "UNLIKELY"]:
        f.write(
            f"{label:15} "
            f"{relevance_counts[label]:,}\n"
        )

    f.write("\nEXACT DUPLICATES\n")
    f.write("-" * 60 + "\n")

    duplicate_groups = [
        paths
        for paths in hashes.values()
        if len(paths) > 1
    ]

    f.write(
        f"Duplicate groups: "
        f"{len(duplicate_groups):,}\n"
    )

    f.write("\nTOP NOVEL CANDIDATES\n")
    f.write("-" * 60 + "\n")

    for record in candidates[:100]:

        f.write(
            f"\n[{record['relevance']}] "
            f"Score {record['relevance_score']}\n"
        )

        f.write(
            f"{record['path']}\n"
        )

        if record["text_preview"]:
            f.write(
                f"Preview: "
                f"{record['text_preview'][:300]}\n"
            )


# ============================================================
# DONE
# ============================================================

print("\n" + "=" * 60)
print("DONE")
print("=" * 60)

print(f"\nResults saved to:")
print(f"  {OUTPUT}")

print("\nImportant files:")
print(f"  {inventory_file.name}")
print(f"  {candidate_file.name}")
print(f"  {summary_file.name}")

print("\nNothing in the original archive was modified.")